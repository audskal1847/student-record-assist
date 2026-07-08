import streamlit as st
import google.generativeai as genai
import PyPDF2
import os
import glob
import requests
import re
import json
from collections import Counter
from docx import Document
from io import BytesIO
from bs4 import BeautifulSoup

# ==========================================================
# 💡 [반영 3] konlpy(형태소 분석기) import - 없으면 자동 폴백
# ==========================================================
KONLPY_AVAILABLE = False
try:
    from konlpy.tag import Okt
    _okt = Okt()
    KONLPY_AVAILABLE = True
except Exception:
    KONLPY_AVAILABLE = False

st.set_page_config(page_title="audskal의 학교생활기록부 분석", layout="wide")
st.title("🏫 객관적이고 체계적인 학생부 분석")
st.markdown("API 키에 맞는 최적의 AI 모델을 자동으로 찾아내어 생기부를 체계적으로 분석합니다.")

# ==========================================================
# 💡 [반영 2] 불용어(제외 단어) 사전 대폭 확장
#   1) 일반 학교/수업/생활 용어  2) 교과명  3) 역량 무관 단순 명사
# ==========================================================
STOPWORDS = set([
    # --- 일반 학교/수업/생활/평가 용어 ---
    "학교", "수업", "활동", "시간", "학생", "선생님", "교사", "담임", "학년", "학기",
    "발표", "참여", "태도", "모습", "친구", "급우", "교실", "학급", "반장", "부반장",
    "역할", "진행", "과정", "결과", "내용", "부분", "생각", "이해", "설명", "질문",
    "대답", "문제", "해결", "노력", "관심", "흥미", "능력", "역량", "자세", "습관",
    "성실", "성실성", "책임감", "리더십", "협력", "협동", "소통", "배려", "존중", "칭찬",
    "선정", "우수", "최선", "적극", "열정", "성장", "발전", "향상", "성취", "목표",
    "계획", "실천", "완성", "수행", "제출", "작성", "정리", "조사", "탐구", "탐색",
    "보고서", "발표회", "토론", "토의", "모둠", "조별", "팀", "친화력", "인성",
    "학업", "성적", "평가", "시험", "과제", "숙제", "수상", "대회", "행사", "축제",
    "동아리", "봉사", "봉사활동", "진로", "진로활동", "자율", "자율활동", "특기",
    "사항", "관련", "여러", "다양", "매우", "항상", "특히", "통해", "위해", "대한",
    "가지", "경우", "때문", "이후", "이전", "당시", "현재", "미래", "자신", "본인",
    # --- 확장: 학교생활기록부 상투어/서술어 명사 ---
    "발휘", "기여", "함양", "습득", "고취", "제고", "증진", "강화", "형성", "확립",
    "기반", "바탕", "중심", "위주", "기초", "기본", "전반", "전체", "모두", "다수",
    "일부", "관찰", "확인", "발견", "제안", "제시", "언급", "표현", "전달", "공유",
    "적용", "활용", "이용", "사용", "실시", "운영", "구성", "포함", "선택", "결정",
    "판단", "분석", "종합", "요약", "설정", "부여", "지도", "안내", "격려", "독려",
    "지속", "꾸준", "성격", "장점", "단점", "특성", "특징", "성향", "면모", "부분적",
    "우수성", "능동", "능동적", "자기", "주도", "주도적", "적극성", "긍정", "긍정적",
    "학습", "공부", "예습", "복습", "질의", "응답", "경청", "메모", "필기", "요점",
    "학교생활", "학교행사", "교내", "교외", "일과", "일상", "생활", "규칙", "질서",
    "친구들", "선생님들", "동급생", "교우", "교우관계", "인간관계", "관계",
    "학기말", "학기초", "연초", "연말", "매일", "매주", "매번", "여름", "겨울",
    "오전", "오후", "방과", "방과후", "점심", "아침", "저녁", "하루", "주말",
    "정도", "수준", "만큼", "이상", "이하", "미만", "초과", "약간", "다소", "상당",
    "이때", "그때", "이번", "다음", "지난", "최근", "차후", "향후", "추후",
    "본격", "실제", "직접", "간접", "구체", "구체적", "명확", "명확히", "분명",
    "학기별", "학년별", "월별", "주별", "단계", "단계별", "차원", "측면", "방향",
    "방법", "방식", "형태", "형식", "종류", "유형", "범위", "영역", "분야",
    "과목", "교과", "교과목", "단원", "차시", "수행평가", "지필", "지필평가",
    # --- 교과명 ---
    "국어", "영어", "수학", "과학", "사회", "역사", "한국사", "세계사", "지리",
    "물리", "화학", "생물", "생명", "지구과학", "윤리", "도덕", "정치", "경제",
    "법과정치", "미술", "음악", "체육", "기술", "가정", "정보", "한문", "제2외국어",
    "문학", "독서", "화법", "작문", "언어", "매체", "확률", "통계", "미적분", "기하",
    "물리학", "화학과", "생명과학", "통합과학", "통합사회", "탐구영역", "수리",
    "국사", "지구", "생활과윤리", "윤리와사상", "동아시아사", "세계지리", "한국지리",
    "정치와법", "사회문화", "경제학", "일본어", "중국어", "독일어", "프랑스어",
    "스페인어", "체육과", "예술", "실용", "심화", "고전", "논술", "수학과", "영어과",
    "국어과", "사회과", "과학과",
])


@st.cache_data(show_spinner=False)
def load_reference_pdfs(pdf_list):
    text = ""
    for pdf_file in pdf_list:
        with open(pdf_file, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    return text


# ==========================================================
# 💡 [반영 3] 형태소 분석기 기반 명사 추출 + [반영 2] 불용어 필터
#   - konlpy 사용 가능 시: Okt로 '명사'만 정확히 추출
#   - 사용 불가 시: 정규식으로 폴백
# ==========================================================
def extract_candidate_keywords(text, top_n=80):
    """
    생기부 텍스트에서 명사형 키워드 후보를 추출하고,
    불용어(학교/교과명/단순명사)를 제거한 뒤 빈도순 후보 목록 반환.
    """
    nouns = []

    if KONLPY_AVAILABLE:
        # ✅ 형태소 분석기(Okt)로 명사만 정확 추출
        try:
            raw_nouns = _okt.nouns(text)
            # 2글자 이상 한글 명사만 사용 (1글자 명사는 노이즈 많음)
            nouns = [n for n in raw_nouns if len(n) >= 2 and re.fullmatch(r'[가-힣]+', n)]
        except Exception:
            nouns = re.findall(r'[가-힣]{2,}', text)
    else:
        # 🔁 폴백: 정규식으로 한글 2글자 이상 단어 추출
        nouns = re.findall(r'[가-힣]{2,}', text)

    filtered = []
    for w in nouns:
        if w in STOPWORDS:
            continue
        # 조사/어미가 붙은 흔한 패턴 간이 제거 (정규식 폴백 시 특히 유효)
        if w.endswith(("하는", "되는", "이라", "에서", "으로", "에게", "까지",
                       "부터", "라고", "면서", "지만", "는데", "다는", "라는")):
            continue
        filtered.append(w)

    counter = Counter(filtered)
    candidates = counter.most_common(top_n)
    return candidates  # [(단어, 빈도), ...]


with st.sidebar:
    st.header("🔑 기본 설정")

    api_provider = st.radio("🤖 API 제공자 선택", ["Google AI Studio", "OpenRouter"])
    api_key = st.text_input("🔑 API 키를 입력하세요", type="password")

    if api_provider == "OpenRouter":
        or_model = st.selectbox("사용할 모델 선택", ["openai/gpt-4o", "anthropic/claude-3.5-sonnet", "google/gemini-2.5-flash"])
        st.markdown("[🔗 OpenRouter 무료 API 키 발급](https://openrouter.ai/keys)")
    else:
        st.markdown("[👉 Google AI Studio 무료 API 키 발급](https://aistudio.google.com/app/apikey)")

    st.markdown("---")

    # ==========================================================
    # 💡 [반영 1] 키워드 분석 On/Off 옵션
    # ==========================================================
    st.subheader("📊 부가 분석 옵션")
    enable_keyword_stats = st.checkbox("핵심 역량 키워드 통계(Top 10) 분석", value=True)
    st.caption(
        "체크 시 AI 호출이 1회 추가되어 시간이 조금 더 소요됩니다.\n\n"
        + ("✅ 형태소 분석기(konlpy) 사용 중 - 명사 추출 정확도 높음"
           if KONLPY_AVAILABLE else
           "ℹ️ konlpy 미설치 - 정규식 방식으로 동작(정상)")
    )

    st.markdown("---")
    st.subheader("📚 내장된 기본 평가 기준 파일")
    pdf_files = glob.glob("*.pdf")
    if pdf_files:
        for f in pdf_files:
            st.write(f"- {f}")
    else:
        st.error("폴더에 기준 PDF 파일이 없습니다!")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. 학교생활기록부 데이터 입력")
    st.info("💡 나이스(NEIS) 원본 PDF는 보안상 안 읽히는 경우가 많습니다. 가급적 아래 빈칸에 내용을 직접 긁어서 붙여넣으세요!")

    student_file = st.file_uploader("📂 학생 생기부 파일 (PDF) 업로드", type=["pdf"], key="student_upload")
    st.markdown("**-- 또는 --**")
    student_text_input = st.text_area("📝 생기부 내용 직접 붙여넣기 (추천)", height=250)

with col2:
    st.subheader("2. 분석 옵션 및 추가 데이터 입력")
    teacher_context = st.text_area(
        "💡 특이사항 및 희망 전공 (예: 생명공학과 진학 희망)",
        height=70
    )

    st.markdown("**🎯 목표 대학 전형 / 전공 가이드북 (선택)**")
    st.info("해당 대학의 가이드북을 업로드하면 평가 기준을 벤치마킹합니다. (※ 결과물에 특정 대학명은 노출되지 않습니다.)")
    univ_guide_file = st.file_uploader("🏫 대학 가이드북 PDF 업로드", type=["pdf"], key="univ_guide_upload")

    st.markdown("**📚 맞춤형 추천 도서 참고 자료 (선택)**")
    default_url = "https://nojaesu.com/category/DIRECTORY/%EA%B5%90%EA%B3%BC%EC%97%B0%EA%B3%84%26%EC%A0%84%EA%B3%B5%EC%A0%81%ED%95%A9%EC%84%9C%20%EA%B8%B0%EC%82%AC%20%EB%AA%A8%EC%9D%8C"
    book_url = st.text_input("🌐 추천 도서 웹사이트 주소(URL)", value=default_url)

    submit_btn = st.button("↵ 🚀 심층 분석 시작", type="primary", use_container_width=True)

st.markdown("---")


def create_word_file(text, keyword_stats=None):
    doc = Document()
    doc.add_heading('AI 생기부 분석 결과 보고서', 0)

    # 💡 키워드 통계 표를 워드 문서 상단에 추가
    if keyword_stats:
        doc.add_heading('핵심 역량 키워드 통계 (Top 10)', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = '순위'
        hdr[1].text = '키워드'
        hdr[2].text = '빈도(회)'
        hdr[3].text = '진로/학과 연관성'
        for item in keyword_stats:
            row = table.add_row().cells
            row[0].text = str(item.get("rank", ""))
            row[1].text = str(item.get("keyword", ""))
            row[2].text = str(item.get("count", ""))
            row[3].text = str(item.get("reason", ""))
        doc.add_paragraph("")

    doc.add_heading('상세 분석 내용', level=1)
    doc.add_paragraph(text)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


if submit_btn:
    if not api_key:
        st.error("왼쪽에 API 키를 먼저 입력해 주세요!")
    elif not pdf_files:
        st.error("기준이 될 PDF 파일이 폴더에 없습니다!")
    elif not student_file and not student_text_input.strip():
        st.error("학생의 생기부 파일(PDF)을 업로드하거나 텍스트를 직접 붙여넣어 주세요!")
    else:
        status_box = st.empty()

        try:
            status_box.info("⏳ [진행상황 1/6] 내장된 기본 가이드북을 학습하는 중입니다...")
            reference_text = load_reference_pdfs(pdf_files)

            univ_guide_text = ""
            if univ_guide_file:
                status_box.info("🏫 [진행상황 2/6] 업로드된 목표 대학 가이드북의 평가 기준을 분석 중입니다...")
                univ_pdf_reader = PyPDF2.PdfReader(univ_guide_file)
                for page in univ_pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        univ_guide_text += extracted + "\n"
            else:
                status_box.info("🏫 [진행상황 2/6] 목표 대학 가이드북이 생략되었습니다. 기본 범용 기준으로 진행합니다.")

            status_box.info("⏳ [진행상황 3/6] 학생의 생기부 데이터를 추출하는 중입니다...")
            student_data_text = ""

            if student_text_input.strip():
                student_data_text = student_text_input
            elif student_file:
                student_pdf_reader = PyPDF2.PdfReader(student_file)
                for page in student_pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        student_data_text += text + "\n"

            if not student_data_text.strip():
                raise Exception("생기부에서 글씨를 읽을 수 없습니다! PDF 대신 빈칸에 직접 붙여넣어 주세요.")

            # ==========================================================
            # 💡 [반영 1] 옵션이 켜진 경우에만 후보 키워드 추출
            # ==========================================================
            candidate_keywords = []
            candidate_str = "후보 없음"
            if enable_keyword_stats:
                engine_name = "형태소 분석기(konlpy)" if KONLPY_AVAILABLE else "정규식 기반"
                status_box.info(f"🔠 [진행상황 4/6] {engine_name}로 핵심 키워드 후보를 추출하고 불용어를 제거하는 중입니다...")
                candidate_keywords = extract_candidate_keywords(student_data_text, top_n=80)
                candidate_str = ", ".join([f"{w}({c})" for w, c in candidate_keywords]) if candidate_keywords else "후보 없음"
            else:
                status_box.info("🔠 [진행상황 4/6] 키워드 통계 분석이 비활성화되어 건너뜁니다.")

            status_box.info("📚 [진행상황 5/6] 추천 도서 목록 및 상세 본문을 수집하는 중입니다...")
            actual_book_data = ""

            if book_url.strip():
                try:
                    headers = {'User-Agent': 'Mozilla/5.0'}
                    response = requests.get(book_url.strip(), headers=headers)
                    response.raise_for_status()
                    soup = BeautifulSoup(response.text, 'html.parser')
                    actual_book_data += soup.get_text(separator=' ', strip=True) + "\n\n"

                    base_url = "/".join(book_url.split("/")[:3])
                    article_links = []
                    for a in soup.find_all('a', href=True):
                        href = a['href']
                        if re.match(r'^/[0-9]+(\?.*)?$', href) or "/entry/" in href:
                            full_url = base_url + href.split('?')[0]
                            if full_url not in article_links:
                                article_links.append(full_url)

                    if article_links:
                        status_box.info(f"📚 [도서 연동] {len(article_links[:10])}개의 구체적인 도서 상세 설명을 추가로 수집 중입니다...")
                        for link in article_links[:10]:
                            try:
                                sub_res = requests.get(link, headers=headers, timeout=5)
                                sub_soup = BeautifulSoup(sub_res.text, 'html.parser')
                                content_area = sub_soup.find('div', class_='entry-content') or sub_soup.find('div', class_='article_view') or sub_soup.body
                                if content_area:
                                    actual_book_data += content_area.get_text(separator=' ', strip=True) + "\n\n"
                            except:
                                pass
                except Exception as e:
                    st.warning(f"⚠️ 입력하신 링크에 접속할 수 없습니다. (오류 메시지: {e})")

            if actual_book_data:
                actual_book_data = actual_book_data.replace("'쌤과 함께! 교과 연계 적합書]", "")
                actual_book_data = actual_book_data.replace("쌤과 함께! 교과 연계 적합書", "")
                actual_book_data = re.sub(r'[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]', '', actual_book_data)

            book_instruction = ""
            if actual_book_data.strip():
                book_instruction = "반드시 아래 제공된 [추천 도서 참고 자료]의 텍스트 안에 '실제로 존재하는 책 제목과 저자'만 추출해서 추천하세요. 자료 안에 적합한 책이 없다면 억지로 지어내지 마세요."
            else:
                book_instruction = "별도로 제공된 도서 목록이 없으므로, AI가 자체적으로 학습한 실존하는 전공 적합 우수 도서를 추천해 주세요. (할루시네이션 절대 금지)"

            status_box.warning(f"🔍 [진행상황 6/6] {api_provider} API를 통해 심층 분석을 시작합니다...")

            # ==========================================================
            # 💡 키워드 통계 전용 프롬프트 (JSON 반환)
            # ==========================================================
            keyword_prompt = f"""
            당신은 대한민국 입학사정관입니다. 아래는 한 학생의 학교생활기록부 원문과,
            프로그램이 1차로 자동 추출한 '명사형 키워드 후보 목록(단어와 등장 횟수)'입니다.

            [학생 생기부 원문]
            {student_data_text}

            [1차 자동 추출 키워드 후보 (단어(빈도) 형식)]
            {candidate_str}

            [학생 희망 전공/특이사항]
            {teacher_context if teacher_context else "특별한 지시사항 없음."}

            🚨 [키워드 선정 규칙 - 매우 중요] 🚨
            1. 아래 유형의 단어는 순위 집계에서 '무조건 제외'하세요.
               - 일반 학교/수업 용어(예: 학교, 수업, 활동, 시간, 학생, 발표, 태도, 노력, 참여 등)
               - 교과명(예: 국어, 영어, 수학, 과학, 사회, 물리, 화학 등)
               - 학생의 역량/전공과 무관한 단순 일반 명사
            2. 반드시 '학생의 진로희망계열 및 대학 학과와 연관성이 있는 전문 키워드',
               또는 '학생의 개별적 역량·전문성·탐구 주제를 효과적으로 드러내는 키워드'만 선정하세요.
            3. 후보 목록의 빈도를 참고하되, 의미가 유사한 단어(예: '유전자', '유전')는 대표 키워드로 통합하고 빈도를 합산해도 됩니다.
            4. 최종적으로 가치 있는 키워드를 '빈도 및 중요도 순'으로 정확히 10개 선정하세요. (10개 미만이면 있는 만큼만)

            🚨 [출력 형식 - 반드시 준수] 🚨
            다른 설명이나 문장 없이, 오직 아래 형식의 JSON 배열만 출력하세요.
            [
              {{"rank": 1, "keyword": "키워드", "count": 빈도수(정수), "reason": "진로/학과와의 연관성을 15자 내외로 간결하게"}}
            ]
            """

            # --- 메인 심층 분석 프롬프트 ---
            prompt = f"""
            당신은 20년 경력의 대한민국 최고 수석 진학 상담 교사이자 입학사정관입니다.

            [담당 교사의 특별 지시사항 및 희망 전공]
            {teacher_context if teacher_context else "특별한 지시사항 없음."}

            [추천 도서 참고 자료 (게시물 상세 본문 포함)]
            {actual_book_data if actual_book_data else "제공된 목록 없음."}

            [기본 범용 대학 평가 기준 자료]
            {reference_text}

            [목표 대학 전형/전공 가이드북 평가 기준 (선택 사항)]
            {univ_guide_text if univ_guide_text else "제공된 목표 대학 가이드북 없음. 기본 범용 평가 기준만 적용할 것."}

            [업로드된 학생의 생기부 내용 (100% 팩트)]
            {student_data_text}

            🚨 [분석의 깊이 및 톤앤매너 (매우 중요)] 🚨
            1. 객관적이고 현실적인 평가 (과장 금지): 학생의 역량을 지나치게 긍정적으로 포장하거나 미사여구('탁월함', '압도적', '완벽함', '뛰어남' 등)를 남발하는 것을 엄격히 금지합니다. 실제 서류평가를 진행하는 냉철하고 객관적인 입학사정관의 시각에서, 학생이 '실제로 수행한 수준'에 맞게 현실적이고 담백하게 서술하세요.
            2. 단순 요약 금지: 단순히 생기부 내용을 요약하거나 줄거리처럼 나열(Summary)하는 것을 금지합니다. 사실(Fact)을 바탕으로 건조하지만 깊이 있는 평가(Analysis)를 작성하세요.

            🚨 [절대 엄수 - 팩트 체크 및 소설 작성 금지 규칙!] 🚨
            1. 학생부 팩트 기반: 업로드된 내용에 없는 과목이나 활동은 단 한 글자도 지어내지 마세요.
            2. 🚫 [진행 중인 학년 기록 부재 지적 절대 금지!]: 최신 학년의 기록이 없는 것은 당연하므로 "기록이 부족함", "학년별 연계가 끊어짐" 등의 지적을 절대 금지합니다.
            3. 도서 추천 규칙: {book_instruction}
            4. 특정 대학명 노출 절대 금지: 동국대학교 등 대학 이름이 직접 등장하면 안 됩니다.
            5. 특정 문구 출력 금지: "'쌤과 함께! 교과 연계 적합書]"와 같은 불필요한 기호는 절대 출력하지 마세요.

            🚨 [절대 엄수 - 출력 형식 및 '출처 100% 일치' 규칙! (가장 중요!)] 🚨
            1. 출처 사전 확정 및 압축: 한 문단에 들어갈 핵심 활동(출처)을 **정확히 2~3개만 먼저 확정**하세요. 학년과 과목은 개별 분리하세요. (예: [1, 2학년 진로] ❌ -> [1학년 진로활동, 2학년 진로활동] ⭕)

            2. 🚫 [문장별 개별 꼬리표 절대 누락 금지!]:
               - 반드시 문단 시작 부분에 **'■ 테마 요약 소제목 [출처1, 출처2]'** 형태로 적으세요.
               - 본문을 작성할 때, 학생의 활동을 언급하는 **모든 문장의 끝에는 반드시 해당 활동의 개별 출처 꼬리표(예: [1학년 진로활동])를 달아야 합니다.**
               - 소제목 옆에 선언한 전체 출처 목록과, 본문 문장 끝에 달린 개별 꼬리표들의 집합은 100% 일치해야 합니다. 즉흥적 추가나 누락을 금지합니다.

            🚨 [항목별 세부 작성 양식 (반드시 지킬 것!)] 🚨
            - 2번 항목(약점 분석): 최소 2가지 이상의 약점을 분석하세요. 약점을 지적할 때는 단순히 약점만 나열하지 말고, **반드시 학생부에 기재된 특정 활동 내용(출처 꼬리표 포함)을 먼저 구체적으로 언급한 뒤**, 이를 보완하기 위한 '구체적인 후속 활동이나 심화 탐구 방안'을 명확하게 제안하세요.

            - 4번 항목(추천 도서 및 연계 활동):
              1) 도서를 나열할 때 반드시 **"■ 교과명(또는 영역명):"** 이라는 굵은 소제목으로 과목/영역을 확실히 구분하여 먼저 적으세요.
              2) 각 교과명 아래에 도서를 1번부터 차례대로 순차 번호를 매겨 나열하세요.
              3) 도서 1개당 반드시 아래의 포맷을 무조건 지키세요.
                 [번호]. <실제 책제목> (저자명 저) - 세부 본문 내용을 바탕으로 한 구체적인 도서 소개 및 추천 이유
                 **연계 활동:** (반드시 줄을 바꾸고 이 굵은 제목을 달 것) 학생의 [O학년 OO과목] 내용과 매칭하여 구체적으로 어떤 심화 탐구/보고서 작성을 할 수 있는지 제안함.

            ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            🚨🚨🚨 [절대 엄수 - 섹션별 '문장 종결 어미' 분리 규칙! (오류 방지 핵심!)] 🚨🚨🚨
            ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            ※ 섹션마다 문장을 끝내는 방식(종결 어미)이 다릅니다. 절대 혼동하지 마세요.

            ▶ [1번, 2번, 4번, 5번 항목] → 반드시 '개조식(명사형 종결)' 사용
               - 모든 문장의 끝을 반드시 '~함', '~임', '~됨', '~판단됨', '~필요함', '~보임', '~드러남' 등 명사형으로 종결할 것.
               - '~합니다', '~입니다', '~됩니다', '~습니다', '~이다', '~한다' 같은 경어체/평서문 종결어미는 절대 사용 금지.
               - (특히 4번 도서 소개 문장과 5번 종합 의견은 길어져도 끝까지 개조식을 유지할 것. 경어체로 풀어지지 말 것!)

            ▶ [3번 항목 중 '면접 예상 질문'] → 반드시 '완결된 질문 문장(의문문/평서문)' 사용
               - 면접관이 학생에게 직접 묻는 '실제 질문'이므로 개조식(~함/~임)을 절대 사용하지 말 것.
               - 반드시 '~까?', '~는가?', '~인가?', '~설명해 보세요.', '~말해 보시오.' 형태의 완전한 문장으로 종결할 것.

            ▶ [3번 항목 중 '추천 심화 탐구 주제'] → 명사구 또는 개조식으로 작성(무방함)
            ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

            위의 모든 규칙과 템플릿을 완벽히 적용하여, 아래 5가지 양식에 맞추어 최종 결과물을 작성해 주세요.
            ### 1. 전공 적합성 및 주요 경쟁력 (테마별 엄선, 객관적/현실적 분석 서술, 100% 일치하는 분리 출처, 개조식 종결)
            ### 2. 범용 평가 기준에 비추어 볼 때 보완이 필요한 약점 (특정 활동 선 언급 후 구체적 보완책 제시, 100% 일치하는 분리 출처, 개조식 종결) ※ 부재중인 학년의 기록 부족 지적 불가!
            ### 3. 추천 심화 탐구 주제 및 면접 예상 질문 3가지 (※ 면접 질문은 반드시 '~다/~까/~요/~오' 형태의 완결된 질문 문장으로 종결!)
            ### 4. 맞춤형 추천 도서 및 연계 활동 제안 (교과명 분류, 도서명(저자) 형식, **연계 활동:** 명시적 서술 필수, 개조식 종결!)
            ### 5. 종합 의견 및 향후 발전 방향 (※ 반드시 개조식 종결! '~합니다/~됩니다' 경어체 절대 금지!)
            """

            # ==========================================================
            # 💡 AI 호출 공통 함수 (제공자 분기)
            # ==========================================================
            def call_ai(prompt_text):
                if api_provider == "Google AI Studio":
                    genai.configure(api_key=api_key)
                    best_model_name = ""
                    for m in genai.list_models():
                        if 'generateContent' in m.supported_generation_methods:
                            best_model_name = m.name.replace("models/", "")
                            if 'flash' in best_model_name or 'pro' in best_model_name:
                                break
                    if best_model_name == "":
                        raise Exception("Google AI Studio에서 사용할 수 있는 모델이 없습니다.")
                    model = genai.GenerativeModel(best_model_name)
                    response = model.generate_content(prompt_text)
                    return response.text

                elif api_provider == "OpenRouter":
                    headers = {
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    }
                    data = {
                        "model": or_model,
                        "messages": [{"role": "user", "content": prompt_text}]
                    }
                    res = requests.post("https://openrouter.ai/api/v1/chat/completions",
                                        headers=headers, json=data, timeout=180)
                    res.raise_for_status()
                    return res.json()['choices'][0]['message']['content']
                return ""

            # ==========================================================
            # 💡 [반영 1] 옵션이 켜진 경우에만 키워드 통계 AI 분석 수행
            # ==========================================================
            keyword_stats = []
            if enable_keyword_stats:
                try:
                    keyword_raw = call_ai(keyword_prompt)
                    json_match = re.search(r'\[.*\]', keyword_raw, re.DOTALL)
                    if json_match:
                        keyword_stats = json.loads(json_match.group())
                except Exception as ke:
                    st.warning(f"⚠️ 키워드 통계 AI 분석 중 오류가 발생하여 로컬 집계로 대체합니다. ({ke})")
                    keyword_stats = [
                        {"rank": i + 1, "keyword": w, "count": c, "reason": "-"}
                        for i, (w, c) in enumerate(candidate_keywords[:10])
                    ]

            # --- 메인 심층 분석 실행 ---
            result_text = call_ai(prompt)

            status_box.success("✅ [분석 완료!] 심층 분석이 완료되었습니다. 결과물을 확인해 주세요!")

            # ==========================================================
            # 💡 키워드 통계 결과 화면 출력
            # ==========================================================
            if enable_keyword_stats and keyword_stats:
                st.subheader("📊 핵심 역량 키워드 통계 (Top 10)")
                engine_badge = "형태소 분석기(konlpy)" if KONLPY_AVAILABLE else "정규식 방식"
                st.caption(
                    f"※ 추출 엔진: {engine_badge} | 학교/수업 등 일반 용어, 교과명, 단순 명사는 제외하고 "
                    "진로·학과·개별 역량 관련 키워드만 집계함."
                )

                table_data = []
                for item in keyword_stats:
                    table_data.append({
                        "순위": item.get("rank", ""),
                        "키워드": item.get("keyword", ""),
                        "빈도(회)": item.get("count", ""),
                        "진로/학과 연관성": item.get("reason", "")
                    })
                st.table(table_data)

                try:
                    chart_data = {str(item.get("keyword", "")): int(item.get("count", 0)) for item in keyword_stats}
                    st.bar_chart(chart_data)
                except Exception:
                    pass

                st.markdown("---")

            st.subheader("📝 상세 분석 결과")
            st.write(result_text)

            word_file = create_word_file(result_text, keyword_stats if enable_keyword_stats else None)
            st.download_button(
                label="📥 분석 결과 워드 다운로드",
                data=word_file,
                file_name="생기부_맞춤형_분석결과.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except requests.exceptions.Timeout:
            status_box.error("오류가 발생했습니다: 데이터 처리에 너무 오랜 시간이 걸려 서버가 응답을 포기했습니다. 데이터 분량을 조금 줄여서 다시 시도해 주세요.")
        except Exception as e:
            status_box.error(f"오류가 발생했습니다: {e}")

st.divider()
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px; font-size: 13px;'>
    🏫 학교생활기록부 분석 시스템 v10.0<br>
    만든이: <b>신선여자고등학교 김명남</b>
</div>
""", unsafe_allow_html=True)
